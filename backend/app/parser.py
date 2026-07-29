import io
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
import logging

logger = logging.getLogger("complaint_system")
logging.basicConfig(level=logging.INFO)

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extracts text from PDF, DOCX, or TXT file bytes.
    Raises ValueError for unsupported or invalid files.
    """
    ext = filename.split(".")[-1].lower()
    
    if ext == "txt":
        logger.info(f"Extracting text from TXT file: {filename}")
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return file_bytes.decode("latin-1")
            except Exception as e:
                raise ValueError(f"Failed to decode text file: {str(e)}")
                
    elif ext == "pdf":
        logger.info(f"Attempting to extract text from PDF file '{filename}' using PyMuPDF (fitz)...")
        text = ""
        pymupdf_success = False
        
        # Try PyMuPDF (fitz)
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                page_text = page.get_text()
                if page_text:
                    text += page_text + "\n"
            doc.close()
            text = text.strip()
            if text:
                logger.info("Successfully extracted text using PyMuPDF.")
                pymupdf_success = True
            else:
                logger.warning("PyMuPDF returned empty text. Attempting fallback...")
        except Exception as e:
            logger.error(f"PyMuPDF extraction failed: {str(e)}. Attempting fallback to pdfplumber...")
            
        # Fallback to pdfplumber
        if not pymupdf_success:
            logger.info("Extracting text from PDF using pdfplumber fallback...")
            try:
                text = ""
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    if len(pdf.pages) == 0:
                        raise ValueError("PDF has 0 pages.")
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                text = text.strip()
                if not text:
                    raise ValueError("No text could be extracted from any page of the PDF.")
                logger.info("Successfully extracted text using pdfplumber fallback.")
            except Exception as e:
                logger.error(f"pdfplumber extraction failed: {str(e)}")
                raise ValueError(f"Failed to parse PDF file. PDF is invalid or contains no extractable text: {str(e)}")
                
        return text
            
    elif ext in ["docx", "doc"]:
        logger.info(f"Extracting text from DOCX file: {filename}")
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            if not text.strip():
                raise ValueError("DOCX contains no extractable text.")
            return text
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")
            
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF, DOCX, and TXT are supported.")

